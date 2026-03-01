#!/usr/bin/env python3
r"""
consolidate_supplementary.py

Fix known data-quality issues and consolidate supplementary assets into
two clean output folders:

  final_sup_table/   – 10 journal-submission Supplementary Tables (ST1–ST10)
  final_sup_data/    – GitHub-only transparency files

Fixes applied (in-memory, non-destructive to originals):
  ST2  → drop 3 fully-empty baseline_macro_auroc columns
  ST4  → add legend note only (no data change; 50% NaN is by-design)
  ST8  → add legend note only (21% NaN is expected for unmappable views)
  ST9  → add legend note only (predicted_zone defined only at K=10%)
  ST11 → keep only regime_validation rows (drop 20 param_sweep_1d rows)

Renumbering (old → new):
  ST1  → ST1   Dataset overview
  ST2  → ST2   Model performance
  ST3  → ST3   DI full table
  ST4  → ST4   Ablation full (XGB + RF)
  ST5  → ST5   Cross-model agreement
  ST6  → ST6   Label permutation
  ST8  → ST7   Pathway detail              (renumbered)
  ST9  → ST8   VAD diagnostics (all K)     (renumbered)
  ST11 → ST9   Simulation summary          (renumbered)
  ST14 → ST10  Unsupervised clustering     (renumbered)

Usage (PowerShell):
  $BASE = "<path-to-outputs>"
  python .\code\compute\13_results\02_consolidate_supplementary_v2.py --outputs-dir $BASE --dry-run
  python .\code\compute\13_results\02_consolidate_supplementary_v2.py --outputs-dir $BASE --clean

"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ═══════════════════════════════════════════════════════════════════════════════
# Configuration
# ═══════════════════════════════════════════════════════════════════════════════

# Journal Supplementary Tables: (old_filename, new_ST_number, short_title, fix_function_name_or_None)
JOURNAL_TABLES: List[Tuple[str, int, str, Optional[str]]] = [
    ("ST1_dataset_overview.csv",           1,  "Dataset overview",              None),
    ("ST2_model_performance.csv",          2,  "Model performance",             "fix_st2"),
    ("ST3_DI_full_table.csv",              3,  "DI full table",                 "fix_st3"),
    ("ST4_ablation_full.csv",              4,  "Ablation full (XGB + RF)",      "fix_st4"),  # legend note only
    ("ST5_cross_model_agreement.csv",      5,  "Cross-model agreement",         None),
    ("ST6_label_permutation_summary.csv",  6,  "Label permutation",             None),
    ("ST8_pathway_detail.csv",             7,  "Pathway detail",                None),       # legend note only
    ("ST9_vad_summary_allK.csv",           8,  "VAD diagnostics (all K)",       None),       # legend note only
    ("ST11_simulation_summary.csv",        9,  "Simulation summary",            "fix_st11"),
    ("ST14_unsupervised_clustering.csv",   10, "Unsupervised clustering",       None),
]

# GitHub-only files: (old_filename, subfolder_in_final_sup_data)
GITHUB_FILES: List[Tuple[str, str]] = [
    ("ST7_shap_stability.csv",                "tables"),
    ("ST12_top_variance_features.csv",        "tables"),
    ("ST13_q4_features.csv",                  "tables"),
    ("ST13_q4_features_view_summary.csv",     "tables"),
    ("ST10_vad_permnull_full.csv",            "tables"),
    ("ST15_context_vs_modality_signal_enrichment.csv", "tables"),
    ("MANIFEST_SUP_TABLE.json",               "provenance"),
]

# Source tables to copy to GitHub provenance folder
SOURCE_TABLE_FILES: List[str] = [
    "ablation_master_rf.csv",
    "ablation_master_xgb.csv",
    "clustering_comparison.csv",
    "fig1_di_curves.csv",
    "fig6_simulation.csv",
    "pc_class_signal.csv",
    "permnull_summary.csv",
    "regime_consensus.csv",
    "signal_enrichment_by_regime.csv",
    "vad_summary.csv",
]

# Legend notes for tables that have expected missingness (written to LEGEND_NOTES.md)
LEGEND_NOTES: Dict[int, str] = {
    4: (
        "ST4 is a union of XGBoost and Random Forest ablation results. "
        "Model-specific columns (e.g. regime_confidence, DI_mean, perf_all_mean) "
        "are intentionally blank for the other model's rows."
    ),
    7: (
        "ST7 (Pathway detail): NA values indicate views excluded from "
        "gene/pathway enrichment analysis due to unmappable feature identifiers."
    ),
    8: (
        "ST8 (VAD diagnostics): The predicted_zone column is reported only at "
        "K = 10%; rows at other K values are NA by design."
    ),
    9: (
        "ST9 (Simulation summary): Contains regime-validation simulation rows only. "
        "Parameter-sweep metadata rows were excluded to avoid misleading missingness."
    ),
}

# Tables to also export as .docx (ST number → formal title for the Word header)
DOCX_TABLES: Dict[int, str] = {
    1: "Dataset and view summary",
    5: "Cross-model agreement in DI-defined regimes",
    6: "Label-permutation null control",
    9: "Simulation regime-recovery validation",
}


# ═══════════════════════════════════════════════════════════════════════════════
# Fix functions
# ═══════════════════════════════════════════════════════════════════════════════

def fix_st2(df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    """Drop 3 fully-empty baseline_macro_auroc columns from ST2."""
    drop_cols = [
        "baseline__macro_auroc__mean",
        "baseline__macro_auroc__ci95_lo",
        "baseline__macro_auroc__ci95_hi",
    ]
    present = [c for c in drop_cols if c in df.columns]
    missing = [c for c in drop_cols if c not in df.columns]

    if not present:
        return df, f"No baseline_auroc columns found (missing={missing})"

    # only drop if all present cols are entirely NA
    if all(df[c].isna().all() for c in present):
        df = df.drop(columns=present)
        note = f"Dropped all-NA baseline_auroc columns: {present}"
        if missing:
            note += f" (not found in file: {missing})"
        return df, note

    # otherwise do NOT drop (protect against accidental data loss)
    non_null = {c: int(df[c].notna().sum()) for c in present}
    return df, f"Did NOT drop baseline_auroc columns (non-null counts={non_null})"


def fix_st3(df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    dot_cols = [c for c in ["DI_pctl_2.5", "DI_pctl_97.5"] if c in df.columns]
    if dot_cols:
        df = df.drop(columns=dot_cols)
        return df, f"Dropped dot-format percentile columns: {dot_cols}"
    return df, "No dot-format percentile columns found"


def fix_st4(df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    if "model" not in df.columns:
        return df, "WARN: 'model' column not found"
    n_filled = int(df["model"].isna().sum())
    if n_filled:
        df["model"] = df["model"].fillna("xgb_bal")
    return df, f"Filled {n_filled} missing 'model' values with 'xgb_bal'"


def fix_st11(df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    """Keep only regime_validation rows (drop param_sweep_1d rows)."""
    # preferred filter
    if "table_source" in df.columns:
        before = len(df)
        df_clean = df[df["table_source"] == "regime_validation"].copy()
        df_clean = df_clean.drop(columns=["table_source"], errors="ignore")
        df_clean = df_clean.dropna(axis=1, how="all")
        after = len(df_clean)
        return df_clean, f"Kept {after} regime_validation rows, dropped {before-after} param_sweep rows."

    # fallback: keep rows that look like simulation (scenario/seed populated)
    if "scenario" in df.columns:
        before = len(df)
        df_clean = df[df["scenario"].notna()].copy()
        after = len(df_clean)
        return df_clean, f"Fallback filter: kept {after} rows with scenario!=NA (dropped {before-after})."

    return df, "WARN: no table_source or scenario column; no filtering applied."


FIX_DISPATCH = {
    "fix_st2": fix_st2,
    "fix_st3": fix_st3,
    "fix_st4": fix_st4,
    "fix_st11": fix_st11,
}


# ═══════════════════════════════════════════════════════════════════════════════
# Word document export
# ═══════════════════════════════════════════════════════════════════════════════

def _set_cell_border(cell, edges: dict):
    """
    Set specific borders on a cell.
    edges: dict of edge_name → (val, size, color)
    Unspecified edges are set to 'none' (invisible).
    """
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "bottom", "left", "right"):
        if edge in edges:
            val, sz, col = edges[edge]
            el = borders.makeelement(
                qn(f"w:{edge}"),
                {qn("w:val"): val, qn("w:sz"): sz, qn("w:color"): col},
            )
        else:
            el = borders.makeelement(
                qn(f"w:{edge}"),
                {qn("w:val"): "none", qn("w:sz"): "0", qn("w:color"): "auto"},
            )
        borders.append(el)
    tc_pr.append(borders)


# Reusable border specs
_RULE = ("single", "8", "000000")   # thin black horizontal rule
_NONE = {}                           # all edges invisible


def _format_cell_value(val) -> str:
    """Format a DataFrame value for display in a Word table cell."""
    if pd.isna(val):
        return ""
    if isinstance(val, float):
        if abs(val) < 1 and val != 0:
            return f"{val:.4f}"
        return f"{val:.2f}" if val != int(val) else str(int(val))
    return str(val)


def write_table_docx(
    df: pd.DataFrame,
    dest_path: Path,
    st_number: int,
    title: str,
    legend_note: Optional[str] = None,
):
    """
    Write a DataFrame as a journal-standard three-rule Word table.

    Three horizontal rules only (classic scientific table style):
      Rule 1 – top of header row
      Rule 2 – bottom of header row (separates header from data)
      Rule 3 – bottom of last data row
    No vertical borders, no shading, no colours.
    """
    if not HAS_DOCX:
        return

    doc = DocxDocument()

    # -- Page setup: landscape if wide tables --------------------------------
    section = doc.sections[0]
    use_landscape = len(df.columns) > 6
    if use_landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # -- Title ---------------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_after = Pt(6)
    run = title_para.add_run(f"Supplementary Table {st_number}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run2 = title_para.add_run(f" \u00b7 {title}")
    run2.font.size = Pt(10)
    run2.font.name = "Arial"

    # -- Table ---------------------------------------------------------------
    ncols = len(df.columns)
    nrows = len(df)
    table = doc.add_table(rows=nrows + 1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    HEADER_FONT_SIZE = Pt(8)
    BODY_FONT_SIZE = Pt(7.5) if ncols > 10 else Pt(8)

    # --- Header row: top rule + bottom rule ---------------------------------
    for j, col_name in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(str(col_name))
        run.bold = True
        run.font.size = HEADER_FONT_SIZE
        run.font.name = "Arial"
        # Rule 1 (top) + Rule 2 (bottom of header)
        _set_cell_border(cell, {"top": _RULE, "bottom": _RULE})

    # --- Data rows: no borders except bottom rule on last row ---------------
    for i in range(nrows):
        is_last = (i == nrows - 1)
        for j in range(ncols):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            val_str = _format_cell_value(df.iloc[i, j])
            run = para.add_run(val_str)
            run.font.size = BODY_FONT_SIZE
            run.font.name = "Arial"
            # Rule 3: bottom of last data row only
            if is_last:
                _set_cell_border(cell, {"bottom": _RULE})
            else:
                _set_cell_border(cell, _NONE)

    # Compact cell padding
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)

    # -- Legend note ----------------------------------------------------------
    if legend_note:
        note_para = doc.add_paragraph()
        note_para.paragraph_format.space_before = Pt(8)
        run = note_para.add_run("Note: ")
        run.bold = True
        run.italic = True
        run.font.size = Pt(8)
        run.font.name = "Arial"
        run2 = note_para.add_run(legend_note)
        run2.italic = True
        run2.font.size = Pt(8)
        run2.font.name = "Arial"

    doc.save(str(dest_path))


# ═══════════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════════

SUP_LABEL_RE = re.compile(r"(?:^|[_\-\s])(?:SUP|ST|SD|S)(\d+)(?:[_\-\s]|$)", flags=re.IGNORECASE)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def find_file(base_dir: Path, filename: str) -> Optional[Path]:
    """Search for filename in base_dir and its subdirectories."""
    # Direct match first
    direct = base_dir / filename
    if direct.exists():
        return direct
    # Search subdirectories (source_tables/, etc.)
    for p in base_dir.rglob(filename):
        return p
    return None


def extract_sup_number_from_filename(p: Path) -> Optional[int]:
    m = SUP_LABEL_RE.search(p.stem)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None

    # also try "S1", "S2" patterns
    m2 = re.search(r"(?:^|[_\-\s])S(\d+)(?:[_\-\s]|$)", p.stem, flags=re.IGNORECASE)
    if m2:
        try:
            return int(m2.group(1))
        except Exception:
            return None

    # handle ST1 / SD1 final packaging names
    m3 = re.search(r"^(?:ST|SD)(\d+)(?:[_\-\s]|$)", p.stem, flags=re.IGNORECASE)
    if m3:
        try:
            return int(m3.group(1))
        except Exception:
            return None

    return None


# ═══════════════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Consolidate and fix supplementary tables/data into final folders."
    )
    parser.add_argument(
        "--outputs-dir", required=True,
        help="Root outputs directory"
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="Preview actions without writing files."
    )
    parser.add_argument(
        "--clean", action="store_true",
        help="Delete final_sup_table/ and final_sup_data/ before writing."
    )
    args = parser.parse_args()

    outputs_dir = Path(args.outputs_dir)
    sup_table_dir = outputs_dir / "results" / "sup_table"
    sup_data_dir = outputs_dir / "results" / "sup_data"
    final_table_dir = outputs_dir / "results" / "final_sup_table"
    final_data_dir = outputs_dir / "results" / "final_sup_data"

    if not args.dry_run and args.clean:
        if final_table_dir.exists():
            shutil.rmtree(final_table_dir)
        if final_data_dir.exists():
            shutil.rmtree(final_data_dir)

    # Validate input
    if not sup_table_dir.exists():
        print(f"[ERROR] sup_table directory not found: {sup_table_dir}", file=sys.stderr)
        sys.exit(1)

    print("=" * 90)
    print("SUPPLEMENTARY CONSOLIDATION")
    print("=" * 90)
    print(f"  Source (tables) : {sup_table_dir}")
    print(f"  Source (data)   : {sup_data_dir}")
    print(f"  Output (journal): {final_table_dir}")
    print(f"  Output (GitHub) : {final_data_dir}")
    print(f"  Dry run         : {args.dry_run}")
    print()

    # ─── Create output directories ──────────────────────────────────────────
    if not args.dry_run:
        final_table_dir.mkdir(parents=True, exist_ok=True)
        final_data_dir.mkdir(parents=True, exist_ok=True)
        (final_data_dir / "tables").mkdir(exist_ok=True)
        (final_data_dir / "provenance").mkdir(exist_ok=True)
        (final_data_dir / "source_tables").mkdir(exist_ok=True)

    # ─── PART 1: Journal Supplementary Tables ───────────────────────────────
    print("─" * 90)
    print("PART 1: JOURNAL SUPPLEMENTARY TABLES (final_sup_table/)")
    print("─" * 90)

    manifest_entries = []
    fix_log = []

    for old_name, new_num, title, fix_fn_name in JOURNAL_TABLES:
        safe_title = re.sub(r"[^a-z0-9]+", "_", title.lower()).strip("_")
        new_name = f"ST{new_num}_{safe_title}.csv"
        src_path = find_file(sup_table_dir, old_name)

        if src_path is None:
            msg = f"  [SKIP] ST{new_num} ({old_name}): SOURCE NOT FOUND"
            print(msg)
            fix_log.append(msg)
            continue

        print(f"\n  ST{new_num:>2} ← {old_name}")
        print(f"        Title : {title}")
        print(f"        Source: {src_path}")
        print(f"        Output: {new_name}")

        # Read
        df = pd.read_csv(src_path)
        rows_before, cols_before = df.shape

        # Apply fix if specified
        fix_note = "No fix needed"
        if fix_fn_name and fix_fn_name in FIX_DISPATCH:
            df, fix_note = FIX_DISPATCH[fix_fn_name](df)
            print(f"        Fix   : {fix_note}")
        elif fix_fn_name:
            print(f"        [WARN] Fix function '{fix_fn_name}' not found; skipping fix.")

        rows_after, cols_after = df.shape
        print(f"        Shape : {rows_before}×{cols_before} → {rows_after}×{cols_after}")

        # Write
        dest_path = final_table_dir / new_name
        docx_name = new_name.replace(".csv", ".docx")
        dest_docx = final_table_dir / docx_name
        if not args.dry_run:
            df.to_csv(dest_path, index=False)
            sha = sha256_file(dest_path)
            # Word document copy (selected tables only)
            if HAS_DOCX and new_num in DOCX_TABLES:
                legend = LEGEND_NOTES.get(new_num, None)
                write_table_docx(df, dest_docx, new_num, DOCX_TABLES[new_num], legend)
                print(f"        DOCX  : {docx_name}")
        else:
            sha = "(dry-run)"

        # Legend note?
        legend = LEGEND_NOTES.get(new_num, None)
        if legend:
            print(f"        Legend : {legend[:80]}...")

        manifest_entries.append({
            "new_ST": f"ST{new_num}",
            "filename": new_name,
            "title": title,
            "old_filename": old_name,
            "rows": rows_after,
            "cols": cols_after,
            "fix_applied": fix_note,
            "legend_note": legend,
            "sha256": sha,
        })

        fix_log.append(
            f"ST{new_num:>2} | {new_name:<55} | {rows_after:>5}×{cols_after:<3} | {fix_note}"
        )

    # ─── PART 2: GitHub-only files ──────────────────────────────────────────
    print("\n" + "─" * 90)
    print("PART 2: GITHUB-ONLY FILES (final_sup_data/)")
    print("─" * 90)

    github_log = []

    for old_name, subfolder in GITHUB_FILES:
        src_path = find_file(sup_table_dir, old_name)
        if src_path is None:
            src_path = find_file(sup_data_dir, old_name) if sup_data_dir.exists() else None
        if src_path is None:
            msg = f"  [SKIP] {old_name}: NOT FOUND"
            print(msg)
            github_log.append(msg)
            continue

        dest_path = final_data_dir / subfolder / old_name
        print(f"  {old_name:<55} → {subfolder}/")
        if not args.dry_run:
            shutil.copy2(src_path, dest_path)
        github_log.append(f"  {old_name:<55} → {subfolder}/")

    # Source tables
    source_tables_src = sup_table_dir / "source_tables"
    if source_tables_src.exists():
        print(f"\n  Copying source_tables/ provenance files...")
        for fname in SOURCE_TABLE_FILES:
            src = source_tables_src / fname
            if src.exists():
                dest = final_data_dir / "source_tables" / fname
                print(f"    {fname}")
                if not args.dry_run:
                    shutil.copy2(src, dest)
                github_log.append(f"    source_tables/{fname}")
            else:
                print(f"    [SKIP] {fname}: not found")

    # Also copy sup_data provenance files
    if sup_data_dir and sup_data_dir.exists():
        for fname in ["MANIFEST_SUP_DATA.json", "SUP_DATA_SUMMARY.md"]:
            src = sup_data_dir / fname
            if src.exists():
                dest = final_data_dir / "provenance" / fname
                print(f"  {fname:<55} → provenance/")
                if not args.dry_run:
                    shutil.copy2(src, dest)

    # ─── PART 3: Write manifest + legend notes + summary ────────────────────
    print("\n" + "─" * 90)
    print("PART 3: WRITING MANIFEST & LEGEND NOTES")
    print("─" * 90)

    if not args.dry_run:
        # Manifest JSON
        manifest = {
            "created_at": datetime.now(timezone.utc).isoformat(),
            "description": "Final consolidated supplementary tables for journal submission",
            "tables": manifest_entries,
            "legend_notes": {f"ST{k}": v for k, v in LEGEND_NOTES.items()},
            "github_files": [g[0] for g in GITHUB_FILES] + SOURCE_TABLE_FILES,
        }
        manifest_path = final_table_dir / "MANIFEST_FINAL.json"
        manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
        print(f"  Written: {manifest_path}")

        # Legend notes markdown
        legend_md_lines = [
            "# Supplementary Table Legend Notes",
            "",
            "These notes should be included in the Supplementary Information section",
            "of the manuscript, either as footnotes to each table or as a combined",
            "legend block.",
            "",
        ]
        for new_num, note in sorted(LEGEND_NOTES.items()):
            entry = next((e for e in manifest_entries if e["new_ST"] == f"ST{new_num}"), None)
            tname = entry["title"] if entry else f"Table {new_num}"
            legend_md_lines.append(f"**Supplementary Table {new_num}** ({tname}): {note}")
            legend_md_lines.append("")

        legend_path = final_table_dir / "LEGEND_NOTES.md"
        legend_path.write_text("\n".join(legend_md_lines), encoding="utf-8")
        print(f"  Written: {legend_path}")

        # Summary markdown
        summary_lines = [
            "# Final Supplementary Tables — Consolidation Summary",
            "",
            f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
            "",
            "## Journal Supplementary Tables (final_sup_table/)",
            "",
            "| ST# | Filename | Title | Rows | Cols | Fix Applied |",
            "|-----|----------|-------|-----:|-----:|-------------|",
        ]
        for e in manifest_entries:
            summary_lines.append(
                f"| {e['new_ST']} | `{e['filename']}` | {e['title']} | "
                f"{e['rows']} | {e['cols']} | {e['fix_applied']} |"
            )
        summary_lines += [
            "",
            "## GitHub-only Files (final_sup_data/)",
            "",
        ]
        for line in github_log:
            summary_lines.append(line)

        summary_path = final_table_dir / "CONSOLIDATION_SUMMARY.md"
        summary_path.write_text("\n".join(summary_lines) + "\n", encoding="utf-8")
        print(f"  Written: {summary_path}")

    # ─── Final summary ──────────────────────────────────────────────────────
    print("\n" + "=" * 90)
    print("CONSOLIDATION COMPLETE")
    print("=" * 90)
    print(f"\n  Journal tables (final_sup_table/): {len(manifest_entries)} files")
    print(f"  GitHub files   (final_sup_data/) : {len(github_log)} files")
    print()
    print("  Fix log:")
    for line in fix_log:
        print(f"    {line}")
    print()

    if args.dry_run:
        print("  *** DRY RUN — no files were written ***")
    else:
        print(f"  Journal folder: {final_table_dir}")
        print(f"  GitHub folder : {final_data_dir}")


if __name__ == "__main__":
    main()